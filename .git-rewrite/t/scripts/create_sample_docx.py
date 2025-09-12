#!/usr/bin/env python3
"""
Simple script to create a sample DOCX resume for testing
"""

import sys
from pathlib import Path

def create_sample_docx():
    """Create a sample DOCX resume using python-docx if available"""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        print("python-docx not available. Install with: pip install python-docx")
        return False
    
    output_path = Path(__file__).parent.parent / "tests" / "sample_resumes" / "lisa_park_resume.docx"
    
    # Create new document
    document = Document()
    
    # Title
    title = document.add_heading('Lisa Park', 0)
    subtitle = document.add_heading('Full-Stack Developer', level=1)
    
    # Contact Info
    contact = document.add_paragraph()
    contact.add_run('Email: lisa.park@email.com | Phone: (555) 654-3210\n')
    contact.add_run('GitHub: github.com/lisapark | LinkedIn: linkedin.com/in/lisapark\n')
    contact.add_run('Location: Los Angeles, CA')
    
    # Professional Summary
    document.add_heading('Professional Summary', level=1)
    summary = document.add_paragraph()
    summary.add_run('Creative full-stack developer with 4 years of experience building ')
    summary.add_run('user-focused applications. Passionate about clean code, great user ')
    summary.add_run('experiences, and bridging the gap between design and engineering. ')
    summary.add_run('Strong background in both frontend and backend technologies.')
    
    # Experience
    document.add_heading('Professional Experience', level=1)
    
    # Spotify
    spotify_title = document.add_paragraph()
    spotify_title.add_run('Full-Stack Developer | Spotify (2020-Present)').bold = True
    
    spotify_duties = document.add_paragraph()
    spotify_duties.add_run('• Built features for artist dashboard used by 50K+ musicians\n')
    spotify_duties.add_run('• Implemented A/B testing framework for UI experiments\n')
    spotify_duties.add_run('• Reduced page load times by 45% through optimization\n')
    spotify_duties.add_run('• Collaborated with design team on user research\n')
    spotify_duties.add_run('• Led frontend architecture decisions for new features')
    
    # Medium
    medium_title = document.add_paragraph()
    medium_title.add_run('Software Developer | Medium (2018-2020)').bold = True
    
    medium_duties = document.add_paragraph()
    medium_duties.add_run('• Developed content management system features\n')
    medium_duties.add_run('• Built recommendation algorithm improving engagement 25%\n')
    medium_duties.add_run('• Mentored 2 junior developers\n')
    medium_duties.add_run('• Led migration from PHP to Node.js')
    
    # Local Agency
    agency_title = document.add_paragraph()
    agency_title.add_run('Junior Developer | Creative Agency (2016-2018)').bold = True
    
    agency_duties = document.add_paragraph()
    agency_duties.add_run('• Developed client websites and e-commerce platforms\n')
    agency_duties.add_run('• Learned foundation of web development\n')
    agency_duties.add_run('• Worked directly with clients on requirements')
    
    # Education
    document.add_heading('Education', level=1)
    education = document.add_paragraph()
    education.add_run('Coding Bootcamp - General Assembly (2016)\n')
    education.add_run('BA Art History - UCLA (2014)')
    
    # Technical Skills
    document.add_heading('Technical Skills', level=1)
    skills = document.add_paragraph()
    skills.add_run('Frontend: React, Vue.js, TypeScript, HTML/CSS, Design Systems\n')
    skills.add_run('Backend: Node.js, Python, PostgreSQL, MongoDB\n')
    skills.add_run('Tools: Git, Docker, AWS, Figma\n')
    skills.add_run('Design: UI/UX Design, User Research, Prototyping')
    
    # Projects
    document.add_heading('Key Projects', level=1)
    projects = document.add_paragraph()
    projects.add_run('Artist Dashboard: Built comprehensive analytics platform for musicians\n')
    projects.add_run('E-commerce Platform: Developed full-stack solution for local businesses\n')
    projects.add_run('Portfolio Websites: Created custom sites for creative professionals')
    
    # Personal
    document.add_heading('Additional Information', level=1)
    personal = document.add_paragraph()
    personal.add_run('Art background brings unique perspective to product development. ')
    personal.add_run('Volunteer teach coding workshops for underrepresented groups. ')
    personal.add_run('Active member of local design and tech communities.')
    
    # Save document
    document.save(str(output_path))
    print(f"Created sample DOCX resume: {output_path}")
    return True

if __name__ == "__main__":
    success = create_sample_docx()
    sys.exit(0 if success else 1)