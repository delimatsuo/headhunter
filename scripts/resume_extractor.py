#!/usr/bin/env python3
"""
Resume Text Extraction Module
Extract text from various resume file formats (PDF, DOCX, images) for LLM analysis
"""

import os
import logging
from typing import Dict, Optional, List
from pathlib import Path
from dataclasses import dataclass
import subprocess

# Built-in libraries for text extraction
import zipfile


@dataclass
class ExtractionResult:
    """Result of text extraction from a resume file"""
    text: str
    success: bool
    file_type: str
    error_message: Optional[str] = None
    metadata: Optional[Dict] = None


class ResumeTextExtractor:
    """Extract text from resume files using built-in Python libraries where possible"""
    
    def __init__(self, log_level: str = "INFO"):
        # Set up logging
        logging.basicConfig(
            level=getattr(logging, log_level),
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        # Supported file types
        self.supported_types = {'.pdf', '.docx', '.doc', '.txt', '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'}
        
        self.logger.info("ResumeTextExtractor initialized")
    
    def detect_file_type(self, file_path: str) -> str:
        """Detect file type from extension"""
        return Path(file_path).suffix.lower()
    
    def extract_text_from_pdf_builtin(self, file_path: str) -> ExtractionResult:
        """Extract text from PDF using built-in tools and pdftotext if available"""
        try:
            # First try pdftotext if available (usually faster and better)
            try:
                result = subprocess.run(
                    ['pdftotext', file_path, '-'],
                    capture_output=True,
                    text=True,
                    check=True
                )
                text = result.stdout
                if text.strip():
                    return ExtractionResult(
                        text=text.strip(),
                        success=True,
                        file_type='pdf',
                        metadata={'method': 'pdftotext'}
                    )
            except (subprocess.CalledProcessError, FileNotFoundError):
                # pdftotext not available or failed, continue to next method
                pass
            
            # Try using PyPDF2 if available
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    
                    if text.strip():
                        return ExtractionResult(
                            text=text.strip(),
                            success=True,
                            file_type='pdf',
                            metadata={'method': 'PyPDF2', 'pages': len(reader.pages)}
                        )
            except ImportError:
                self.logger.debug("PyPDF2 not available")
            except Exception as e:
                self.logger.warning(f"PyPDF2 extraction failed: {e}")
            
            # Fallback message
            return ExtractionResult(
                text="",
                success=False,
                file_type='pdf',
                error_message="PDF text extraction requires pdftotext or PyPDF2. Install with: pip install PyPDF2"
            )
            
        except Exception as e:
            return ExtractionResult(
                text="",
                success=False,
                file_type='pdf',
                error_message=f"PDF extraction failed: {str(e)}"
            )
    
    def extract_text_from_docx_builtin(self, file_path: str) -> ExtractionResult:
        """Extract text from DOCX using built-in tools and python-docx if available"""
        try:
            # Try using python-docx if available
            try:
                from docx import Document
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                    text += "\n"
                
                if text.strip():
                    return ExtractionResult(
                        text=text.strip(),
                        success=True,
                        file_type='docx',
                        metadata={'method': 'python-docx'}
                    )
            except ImportError:
                self.logger.debug("python-docx not available")
            except Exception as e:
                self.logger.warning(f"python-docx extraction failed: {e}")
            
            # Fallback: try manual XML extraction (basic)
            try:
                text = self._extract_docx_xml(file_path)
                if text.strip():
                    return ExtractionResult(
                        text=text.strip(),
                        success=True,
                        file_type='docx',
                        metadata={'method': 'xml_extraction'}
                    )
            except Exception as e:
                self.logger.warning(f"XML extraction failed: {e}")
            
            return ExtractionResult(
                text="",
                success=False,
                file_type='docx',
                error_message="DOCX text extraction requires python-docx. Install with: pip install python-docx"
            )
            
        except Exception as e:
            return ExtractionResult(
                text="",
                success=False,
                file_type='docx',
                error_message=f"DOCX extraction failed: {str(e)}"
            )
    
    def _extract_docx_xml(self, file_path: str) -> str:
        """Extract text from DOCX by parsing the XML manually"""
        try:
            with zipfile.ZipFile(file_path, 'r') as docx:
                # Read the main document XML
                xml_content = docx.read('word/document.xml')
                xml_str = xml_content.decode('utf-8')
                
                # Remove XML tags and extract text
                # This is a simple approach - real DOCX parsing is more complex
                import re
                # Find text between <w:t> tags
                text_pattern = r'<w:t[^>]*>([^<]*)</w:t>'
                matches = re.findall(text_pattern, xml_str)
                
                text = '\n'.join(matches)
                return text
        except Exception as e:
            self.logger.error(f"Manual DOCX extraction failed: {e}")
            return ""
    
    def extract_text_from_image_ocr(self, file_path: str) -> ExtractionResult:
        """Extract text from image using OCR"""
        try:
            # Try tesseract OCR if available
            try:
                result = subprocess.run(
                    ['tesseract', file_path, 'stdout'],
                    capture_output=True,
                    text=True,
                    check=True
                )
                text = result.stdout
                if text.strip():
                    return ExtractionResult(
                        text=text.strip(),
                        success=True,
                        file_type='image',
                        metadata={'method': 'tesseract'}
                    )
            except (subprocess.CalledProcessError, FileNotFoundError):
                pass
            
            # Try with pytesseract if available
            try:
                import pytesseract
                from PIL import Image
                
                image = Image.open(file_path)
                text = pytesseract.image_to_string(image)
                
                if text.strip():
                    return ExtractionResult(
                        text=text.strip(),
                        success=True,
                        file_type='image',
                        metadata={'method': 'pytesseract'}
                    )
            except ImportError:
                self.logger.debug("pytesseract/PIL not available")
            except Exception as e:
                self.logger.warning(f"pytesseract extraction failed: {e}")
            
            return ExtractionResult(
                text="",
                success=False,
                file_type='image',
                error_message="OCR requires tesseract or pytesseract. Install tesseract or: pip install pytesseract pillow"
            )
            
        except Exception as e:
            return ExtractionResult(
                text="",
                success=False,
                file_type='image',
                error_message=f"OCR extraction failed: {str(e)}"
            )
    
    def extract_text_from_file(self, file_path: str) -> ExtractionResult:
        """Extract text from resume file based on file type"""
        if not os.path.exists(file_path):
            return ExtractionResult(
                text="",
                success=False,
                file_type='unknown',
                error_message=f"File not found: {file_path}"
            )
        
        file_type = self.detect_file_type(file_path)
        self.logger.info(f"Extracting text from {file_type} file: {Path(file_path).name}")
        
        if file_type == '.txt':
            # Plain text file
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                return ExtractionResult(
                    text=text.strip(),
                    success=True,
                    file_type='txt',
                    metadata={'method': 'plain_text'}
                )
            except UnicodeDecodeError:
                # Try with different encoding
                try:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        text = f.read()
                    return ExtractionResult(
                        text=text.strip(),
                        success=True,
                        file_type='txt',
                        metadata={'method': 'plain_text', 'encoding': 'latin-1'}
                    )
                except Exception as e:
                    return ExtractionResult(
                        text="",
                        success=False,
                        file_type='txt',
                        error_message=f"Text file encoding error: {str(e)}"
                    )
        
        elif file_type == '.pdf':
            return self.extract_text_from_pdf_builtin(file_path)
        
        elif file_type == '.docx':
            return self.extract_text_from_docx_builtin(file_path)
        
        elif file_type == '.doc':
            # Legacy DOC format - try antiword if available
            try:
                result = subprocess.run(
                    ['antiword', file_path],
                    capture_output=True,
                    text=True,
                    check=True
                )
                text = result.stdout
                return ExtractionResult(
                    text=text.strip(),
                    success=True,
                    file_type='doc',
                    metadata={'method': 'antiword'}
                )
            except (subprocess.CalledProcessError, FileNotFoundError):
                return ExtractionResult(
                    text="",
                    success=False,
                    file_type='doc',
                    error_message="DOC extraction requires antiword. Install with: brew install antiword"
                )
        
        elif file_type in {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'}:
            return self.extract_text_from_image_ocr(file_path)
        
        else:
            return ExtractionResult(
                text="",
                success=False,
                file_type=file_type,
                error_message=f"Unsupported file type: {file_type}. Supported types: {', '.join(self.supported_types)}"
            )
    
    def extract_text_from_multiple_files(self, file_paths: List[str]) -> List[ExtractionResult]:
        """Extract text from multiple resume files"""
        results = []
        for file_path in file_paths:
            result = self.extract_text_from_file(file_path)
            results.append(result)
            
            if result.success:
                self.logger.info(f"Successfully extracted text from {Path(file_path).name}")
            else:
                self.logger.warning(f"Failed to extract text from {Path(file_path).name}: {result.error_message}")
        
        return results
    
    def get_extraction_summary(self, results: List[ExtractionResult]) -> Dict:
        """Get summary statistics for extraction results"""
        total = len(results)
        successful = sum(1 for r in results if r.success)
        failed = total - successful
        
        file_types = {}
        for result in results:
            file_type = result.file_type
            if file_type not in file_types:
                file_types[file_type] = {'total': 0, 'successful': 0}
            file_types[file_type]['total'] += 1
            if result.success:
                file_types[file_type]['successful'] += 1
        
        return {
            'total_files': total,
            'successful': successful,
            'failed': failed,
            'success_rate': (successful / total * 100) if total > 0 else 0,
            'file_types': file_types
        }


def main():
    """CLI interface for resume text extraction"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Resume Text Extraction Tool')
    parser.add_argument('files', nargs='+', help='Resume file paths')
    parser.add_argument('-o', '--output', help='Output directory for extracted text files')
    parser.add_argument('--log-level', default='INFO', 
                       choices=['DEBUG', 'INFO', 'WARNING', 'ERROR'],
                       help='Logging level')
    
    args = parser.parse_args()
    
    # Initialize extractor
    extractor = ResumeTextExtractor(log_level=args.log_level)
    
    # Extract text from files
    results = extractor.extract_text_from_multiple_files(args.files)
    
    # Save extracted text if output directory specified
    if args.output:
        output_dir = Path(args.output)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        for i, result in enumerate(results):
            if result.success:
                input_file = Path(args.files[i])
                output_file = output_dir / f"{input_file.stem}_extracted.txt"
                
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(result.text)
                
                print(f"Extracted text saved to: {output_file}")
    
    # Print summary
    summary = extractor.get_extraction_summary(results)
    print("\nExtraction Summary:")
    print(f"Total Files: {summary['total_files']}")
    print(f"Successful: {summary['successful']}")
    print(f"Failed: {summary['failed']}")
    print(f"Success Rate: {summary['success_rate']:.1f}%")
    
    # Print file type breakdown
    print("\nBy File Type:")
    for file_type, stats in summary['file_types'].items():
        print(f"  {file_type}: {stats['successful']}/{stats['total']} successful")
    
    # Print failed extractions
    failed_results = [r for r in results if not r.success]
    if failed_results:
        print("\nFailed Extractions:")
        for i, result in enumerate(failed_results):
            file_path = [f for j, f in enumerate(args.files) if results[j] == result][0]
            print(f"  {Path(file_path).name}: {result.error_message}")


if __name__ == "__main__":
    main()